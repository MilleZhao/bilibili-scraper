'''
=============================================================================
B站视频 评论+弹幕 综合分析报告
完全基于 build_report.py 的模板重写,全中文,一次跑通
=============================================================================
依赖: pip install python-docx matplotlib openpyxl
运行: py -3.13 build_full_report.py
=============================================================================
'''
import json, re, os
from pathlib import Path
from datetime import datetime, timezone, timedelta
from collections import Counter

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from matplotlib.font_manager import FontProperties

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from openpyxl.cell.cell import ILLEGAL_CHARACTERS_RE

BASE    = Path(__file__).parent / 'outputs'
IMG_DIR = BASE / 'charts'
IMG_DIR.mkdir(parents=True, exist_ok=True)

CMT_SRC  = BASE / 'comments.json'
DM_SRC   = BASE / 'danmaku_dedup.json'
DM_RAW   = BASE / 'danmaku_BV11mFLziEyP.json'
DM_STAGE = BASE / 'danmaku_stage_stats.json'
OUT_DOCX = BASE / 'B站视频综合分析报告.docx'
BVID     = 'BV11mFLziEyP'

# ============================================================
# 中文字体
# ============================================================
try:
    FONT_CN    = FontProperties(fname=r'C:\Windows\Fonts\msyh.ttc', size=10)
    FONT_TITLE = FontProperties(fname=r'C:\Windows\Fonts\msyh.ttc', size=12)
except:
    FONT_CN = FONT_TITLE = None
plt.rcParams['axes.unicode_minus'] = False

# ============================================================
# 加载数据
# ============================================================
print('[1/7] 加载数据...')
comments_raw = json.loads(CMT_SRC.read_text(encoding='utf-8'))
danmaku_dedup = json.loads(DM_SRC.read_text(encoding='utf-8'))
danmaku_raw   = json.loads(DM_RAW.read_text(encoding='utf-8'))
dm_stage      = json.loads(DM_STAGE.read_text(encoding='utf-8'))

for c in comments_raw:
    c['content']  = ILLEGAL_CHARACTERS_RE.sub('', c.get('content', '') or '')
    c['username'] = ILLEGAL_CHARACTERS_RE.sub('', c.get('username', '') or '')
for d in danmaku_raw:
    d['text'] = ILLEGAL_CHARACTERS_RE.sub('', d.get('text', '') or '')

# ============================================================
# 评论清洗 (复用 build_report.py 的逻辑)
# ============================================================
print('[2/7] 清洗评论...')

def is_emoji(ch):
    cp = ord(ch)
    return (0x1F300<=cp<=0x1FAFF or 0x2600<=cp<=0x27BF or cp==0x2764
         or 0x1F900<=cp<=0x1F9FF or 0x1FA00<=cp<=0x1FA6F
         or 0x200D<=cp<=0x200D or 0xFE0F<=cp<=0xFE0F
         or 0x2702<=cp<=0x27B0 or 0x1F600<=cp<=0x1F64F
         or 0x1F680<=cp<=0x1F6FF or 0x1F1E0<=cp<=0x1F1FF)

def has_text(text):
    cleaned = ''.join(ch for ch in text if not is_emoji(ch))
    cleaned = re.sub(r'[\(\（]\s*[^\w\s]{2,}\s*[\)\）]', '', cleaned)
    chinese = len(re.findall(r'[\u4e00-\u9fff]', cleaned))
    english = len(re.findall(r'[a-zA-Z]{3,}', cleaned))
    return chinese >= 2 or english >= 1

def is_ad(text):
    patterns = [
        r'https?://\S+|www\.\S+', r'1[3-9]\d{9}|\d{3}[-\s]?\d{4}[-\s]?\d{4}',
        r'[加➕﹢＋]?\s*(微|v信|威信|VX|vx|WX|wx)[:：]?\s*\w{5,}',
        r'[加➕﹢＋]?\s*(qq|QQ|扣扣|Q|q)[:：]?\s*\d{5,}',
        r'扫码|二维码|QR码', r'免费领取|免费送|低价|优惠券|薅羊毛|返利|刷单|兼职|日赚|代练|代打',
    ]
    return any(re.search(p, text, re.I) for p in patterns)

def is_valid_comment(c):
    text = (c.get('content') or '').strip()
    if not text: return False
    if not has_text(text): return False
    if is_ad(text): return False
    if len(text) <= 2: return False
    if text in ('好','加油','支持','6','厉害','牛','顶','打卡','来了','第一','棒','不错','nb','NB','哈哈','哈哈哈','可以的','好家伙','三连'):
        return False
    if re.fullmatch(r'[\d\s,，]+楼', text): return False
    if len(set(text)) <= 3 and len(text) >= 3: return False
    if re.fullmatch(r'[\d\s\.\,\!\?\#\@\$\%\^\&\*\(\)\+\=~\-—–_/\\|:;]+', text): return False
    return True

comments_valid = [c for c in comments_raw if is_valid_comment(c)]
comments_main  = [c for c in comments_valid if not c.get('is_sub')]
comments_sub   = [c for c in comments_valid if c.get('is_sub')]

cmt_emoji   = sum(1 for c in comments_raw if not has_text((c.get('content') or '').strip()))
cmt_ad      = sum(1 for c in comments_raw if is_ad((c.get('content') or '').strip()))
cmt_invalid = len(comments_raw) - len(comments_valid)
cmt_other   = cmt_invalid - cmt_emoji - cmt_ad

# ============================================================
# 统计指标
# ============================================================
print('[3/7] 计算指标...')

TOTAL_BILI   = 24754
CMT_CRAWLED  = len(comments_raw)
CMT_VALID    = len(comments_valid)
CMT_RET      = CMT_VALID / max(CMT_CRAWLED, 1) * 100

DM_RAW_COUNT  = dm_stage.get('raw', len(danmaku_raw))
DM_DEDUP      = len(danmaku_dedup)
DM_REPEAT_TOT = dm_stage.get('total_repeat', sum(g['count'] for g in danmaku_dedup))
DM_REPEAT_MAX = dm_stage.get('max_repeat', 0)
DM_CLEAN      = dm_stage.get('clean', DM_RAW_COUNT)
DEDUP_RATE    = DM_DEDUP / max(DM_RAW_COUNT, 1) * 100

# 评论日分布
def parse_ts(ts_val):
    if isinstance(ts_val, (int, float)) and ts_val > 1000000000:
        return datetime.fromtimestamp(ts_val, tz=timezone.utc).astimezone().strftime('%Y-%m-%d')
    if isinstance(ts_val, str):
        if ts_val.isdigit() and len(ts_val) >= 10:
            return datetime.fromtimestamp(int(ts_val), tz=timezone.utc).astimezone().strftime('%Y-%m-%d')
        try:
            return datetime.fromisoformat(ts_val).strftime('%Y-%m-%d')
        except:
            pass
    return None

cmt_daily = Counter()
for c in comments_valid:
    day = parse_ts(c.get('timestamp') or c.get('ctime'))
    if day: cmt_daily[day] += 1

dm_daily = Counter()
for d in danmaku_raw:
    day = d.get('date', '')
    if day: dm_daily[day] += 1

# 评论统计
cmt_likes       = [c.get('likes', 0) for c in comments_valid if c.get('likes', 0) > 0]
cmt_lengths     = [len(c.get('content', '')) for c in comments_valid if c.get('content')]
cmt_top_users   = Counter()
for c in comments_valid:
    cmt_top_users[(c.get('username', '') or '未知')] += 1

# 弹幕统计
dm_counts     = [g['count'] for g in danmaku_dedup]
dm_first_ats  = [g['first_at'] for g in danmaku_dedup if g['first_at'] > 0]
VIDEO_DUR     = max(dm_first_ats) if dm_first_ats else 600

# 弹幕高频词（加权）
dm_words = Counter()
for g in danmaku_dedup:
    text = g['text']
    w = g['count']
    for wrd in re.findall(r'[\u4e00-\u9fff]{2,}', text):
        dm_words[wrd] += w

# ============================================================
# 图表
# ============================================================
print('[4/7] 生成图表...')

DPI   = 150
BLUE  = '#2E74B5'
PINK  = '#FB7299'
DARK  = '#1F3A5F'
GREEN = '#27AE60'
PURPLE= '#8E44AD'
GRAY  = '#95A5A6'

# --- 图1: 评论+弹幕每日分布 ---
fig, ax = plt.subplots(figsize=(10, 5))
all_dates = sorted(set(list(cmt_daily.keys()) + list(dm_daily.keys())))
x = range(len(all_dates))
w = 0.35
ax.bar([i - w/2 for i in x], [cmt_daily.get(d, 0) for d in all_dates], w, color=BLUE, alpha=0.85, label='\u8bc4\u8bba')
ax.bar([i + w/2 for i in x], [dm_daily.get(d, 0) for d in all_dates], w, color=PINK, alpha=0.85, label='\u5f39\u5e55')
ax.set_xticks(x)
ax.set_xticklabels(all_dates, rotation=45, ha='right', fontsize=8)
ax.legend(loc='upper right', prop=FONT_CN)
ax.set_title('\u8bc4\u8bba & \u5f39\u5e55 \u6bcf\u65e5\u5206\u5e03', fontproperties=FONT_TITLE)
ax.set_ylabel('\u6570\u91cf', fontproperties=FONT_CN)
ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v:,.0f}'))
ax.grid(axis='y', alpha=0.3)
fig.tight_layout()
fig.savefig(IMG_DIR / 'chart1_daily.png', dpi=DPI)
plt.close()

# --- 图2: 弹幕重复次数排名 Top200 ---
fig, ax = plt.subplots(figsize=(8, 4))
show = sorted(dm_counts, reverse=True)[:200]
ax.bar(range(len(show)), show, color=PINK, alpha=0.8, width=1, edgecolor='none')
ax.set_title('\u5f39\u5e55\u91cd\u590d\u6b21\u6570\u6392\u540d\uff08Top 200\uff09', fontproperties=FONT_TITLE)
ax.set_xlabel('\u6392\u540d', fontproperties=FONT_CN)
ax.set_ylabel('\u91cd\u590d\u6b21\u6570', fontproperties=FONT_CN)
ax.grid(axis='y', alpha=0.3)
for i in [0,1,2,4,9,19,49,99,199]:
    if i < len(show):
        ax.annotate(str(show[i]), (i, show[i]), textcoords='offset points',
                    xytext=(0, 5), fontsize=7, ha='center')
fig.tight_layout()
fig.savefig(IMG_DIR / 'chart2_dm_freq.png', dpi=DPI)
plt.close()

# --- 图3: 弹幕时间密度 ---
fig, ax = plt.subplots(figsize=(10, 3.5))
bins = max(20, min(200, int(VIDEO_DUR // 10)))
ax.hist(dm_first_ats, bins=bins, color=PINK, alpha=0.8, edgecolor='white', linewidth=0.3)
ax.set_title('\u5f39\u5e55\u9996\u6b21\u51fa\u73b0\u65f6\u95f4\u5bc6\u5ea6\uff08\u53bb\u91cd\u540e\uff09', fontproperties=FONT_TITLE)
ax.set_xlabel('\u89c6\u9891\u65f6\u95f4\uff08\u79d2\uff09', fontproperties=FONT_CN)
ax.set_ylabel('\u5f39\u5e55\u6570\u91cf', fontproperties=FONT_CN)
ax.grid(axis='y', alpha=0.3)
fig.tight_layout()
fig.savefig(IMG_DIR / 'chart3_dm_density.png', dpi=DPI)
plt.close()

# --- 图4: 评论点赞分布 ---
fig, ax = plt.subplots(figsize=(8, 4))
if cmt_likes:
    bins_likes = [0,1,5,10,20,50,100,200,500,1000,5000,10000,50000]
    ax.hist(cmt_likes, bins=bins_likes, color=BLUE, alpha=0.8, edgecolor='white')
    ax.set_xscale('log')
    ax.set_xticks([1,10,100,1000,10000,50000])
    ax.set_xticklabels(['1','10','100','1k','10k','50k'])
ax.set_title('\u8bc4\u8bba\u70b9\u8d5e\u6570\u5206\u5e03', fontproperties=FONT_TITLE)
ax.set_ylabel('\u8bc4\u8bba\u6570\u91cf', fontproperties=FONT_CN)
ax.grid(axis='y', alpha=0.3)
fig.tight_layout()
fig.savefig(IMG_DIR / 'chart4_likes.png', dpi=DPI)
plt.close()

# --- 图5: 活跃用户 Top15 ---
fig, ax = plt.subplots(figsize=(8, 5))
top15 = cmt_top_users.most_common(15)
users = [u for u, _ in reversed(top15)]
counts= [c for _, c in reversed(top15)]
ax.barh(range(len(users)), counts, color=BLUE, alpha=0.8)
ax.set_yticks(range(len(users)))
ax.set_yticklabels(users, fontsize=9, fontproperties=FONT_CN)
ax.set_title('\u6700\u6d3b\u8dc3\u8bc4\u8bba\u7528\u6237 Top 15', fontproperties=FONT_TITLE)
ax.set_xlabel('\u8bc4\u8bba\u6570', fontproperties=FONT_CN)
for i, v in enumerate(counts):
    ax.text(v+1, i, str(v), va='center', fontsize=8, fontproperties=FONT_CN)
fig.tight_layout()
fig.savefig(IMG_DIR / 'chart5_top_users.png', dpi=DPI)
plt.close()

# --- 图6: 弹幕高频词 Top20 ---
fig, ax = plt.subplots(figsize=(8, 5))
top20 = dm_words.most_common(20)
words = [w for w, _ in reversed(top20)]
freqs = [c for _, c in reversed(top20)]
ax.barh(range(len(words)), freqs, color=PINK, alpha=0.8)
ax.set_yticks(range(len(words)))
ax.set_yticklabels(words, fontsize=10, fontproperties=FONT_CN)
ax.set_title('\u5f39\u5e55\u9ad8\u9891\u8bcd Top 20\uff08\u52a0\u6743\uff09', fontproperties=FONT_TITLE)
ax.set_xlabel('\u52a0\u6743\u9891\u6b21', fontproperties=FONT_CN)
for i, v in enumerate(freqs):
    ax.text(v+1, i, str(v), va='center', fontsize=8, fontproperties=FONT_CN)
fig.tight_layout()
fig.savefig(IMG_DIR / 'chart6_dm_words.png', dpi=DPI)
plt.close()

# --- 图7: 体量对比 ---
fig, ax = plt.subplots(figsize=(6, 4))
cats = ['\u6709\u6548\u8bc4\u8bba', '\u5f39\u5e55\uff08\u53bb\u91cd\u524d\uff09', '\u5f39\u5e55\uff08\u53bb\u91cd\u540e\uff09']
vals = [CMT_VALID, DM_REPEAT_TOT, DM_DEDUP]
bars = ax.bar(cats, vals, color=[BLUE, PINK, PURPLE], alpha=0.85)
ax.set_title('\u5185\u5bb9\u4f53\u91cf\u5bf9\u6bd4', fontproperties=FONT_TITLE)
ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v:,.0f}'))
for bar, val in zip(bars, vals):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+500,
            f'{val:,}', ha='center', fontsize=10, fontproperties=FONT_CN)
ax.grid(axis='y', alpha=0.3)
fig.tight_layout()
fig.savefig(IMG_DIR / 'chart7_compare.png', dpi=DPI)
plt.close()

print(f'    7\u5f20\u56fe\u8868\u5df2\u4fdd\u5b58\u5230 {IMG_DIR}')

# ============================================================
# Word 报告 (完全复用 build_report.py 的模板)
# ============================================================
print('[5/7] \u751f\u6210Word\u62a5\u544a...')

doc = Document()
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(1.0)
section.right_margin  = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.0)

# 样式
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(6)
style.paragraph_format.line_spacing = 1.10
style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')

for level, (size, color_hex) in enumerate([(16,'2E74B5'),(13,'2E74B5'),(12,'1F4D78')], 1):
    hs = doc.styles[f'Heading {level}']
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor(*tuple(int(color_hex[i:i+2],16) for i in (0,2,4)))
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')

BLUE_H   = RGBColor(0x2E, 0x74, 0xB5)
DARK_B   = RGBColor(0x1F, 0x3A, 0x5F)
INK      = RGBColor(0x0B, 0x25, 0x45)
TBL_HEAD = '2E74B5'
TBL_FILL = 'F4F6F9'

def add_para(text, bold=False, size=None, color=None, alignment=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if alignment is not None: p.alignment = alignment
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    return p

def set_cell(cell, text, bold=False, size=10, color=INK, align='left', fill=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if fill:
        shading = parse_xml('<w:shd ' + nsdecls('w') + ' w:fill=' + chr(34) + fill + chr(34) + '/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        '<w:tcMar ' + nsdecls('w') + '>'
        '<w:top w:w=' + chr(34) + '60' + chr(34) + ' w:type=' + chr(34) + 'dxa' + chr(34) + '/>'
        '<w:bottom w:w=' + chr(34) + '60' + chr(34) + ' w:type=' + chr(34) + 'dxa' + chr(34) + '/>'
        '<w:start w:w=' + chr(34) + '100' + chr(34) + ' w:type=' + chr(34) + 'dxa' + chr(34) + '/>'
        '<w:end w:w=' + chr(34) + '100' + chr(34) + ' w:type=' + chr(34) + 'dxa' + chr(34) + '/>'
        '</w:tcMar>'
    )
    tc_pr.append(tc_mar)

def set_table_borders(table, color='BFBFBF', size='4'):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml('<w:tblPr ' + nsdecls('w') + '/>')
    borders = parse_xml(
        '<w:tblBorders ' + nsdecls('w') + '>'
        '<w:top w:val=' + chr(34) + 'single' + chr(34) + ' w:sz=' + chr(34) + size + chr(34) + ' w:space=' + chr(34) + '0' + chr(34) + ' w:color=' + chr(34) + color + chr(34) + '/>'
        '<w:left w:val=' + chr(34) + 'single' + chr(34) + ' w:sz=' + chr(34) + size + chr(34) + ' w:space=' + chr(34) + '0' + chr(34) + ' w:color=' + chr(34) + color + chr(34) + '/>'
        '<w:bottom w:val=' + chr(34) + 'single' + chr(34) + ' w:sz=' + chr(34) + size + chr(34) + ' w:space=' + chr(34) + '0' + chr(34) + ' w:color=' + chr(34) + color + chr(34) + '/>'
        '<w:right w:val=' + chr(34) + 'single' + chr(34) + ' w:sz=' + chr(34) + size + chr(34) + ' w:space=' + chr(34) + '0' + chr(34) + ' w:color=' + chr(34) + color + chr(34) + '/>'
        '<w:insideH w:val=' + chr(34) + 'single' + chr(34) + ' w:sz=' + chr(34) + size + chr(34) + ' w:space=' + chr(34) + '0' + chr(34) + ' w:color=' + chr(34) + color + chr(34) + '/>'
        '<w:insideV w:val=' + chr(34) + 'single' + chr(34) + ' w:sz=' + chr(34) + size + chr(34) + ' w:space=' + chr(34) + '0' + chr(34) + ' w:color=' + chr(34) + color + chr(34) + '/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_img(path, width=6.3):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.add_paragraph()

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    for ci, h in enumerate(headers):
        set_cell(t.cell(0, ci), h, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), fill=TBL_HEAD, align='center')
    for ri, row_data in enumerate(rows, 1):
        for ci, val in enumerate(row_data):
            is_final = (ri == len(rows))
            is_neg   = str(val).startswith('-')
            c = RGBColor(0x9B,0x1C,0x1C) if is_neg else (DARK_B if is_final else INK)
            fill = 'E8EEF5' if is_final else (TBL_FILL if ri % 2 == 0 else None)
            set_cell(t.cell(ri, ci), str(val), bold=is_final, color=c, fill=fill,
                     align='center' if ci > 0 else 'left')
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return t

# ============================================================
# 封面
# ============================================================
add_para('B站视频 评论+弹幕 综合分析报告', bold=True, size=22, color=DARK_B, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(f'{BVID}  —  当世界过分诚实，我们要如何保持好奇与勇气', size=12, color=BLUE_H, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

p_line = doc.add_paragraph()
p_line.paragraph_format.space_after = Pt(14)
run_line = p_line.add_run('_' * 60)
run_line.font.size = Pt(6)
run_line.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

meta = [
    ('报告日期', datetime.now().strftime('%Y年%m月%d日')),
    ('数据来源', 'bilibili.com — bilibili-api 异步爬虫'),
    ('爬取时间', '2026年6月'),
    ('分析工具', 'Python + Matplotlib + python-docx'),
]
make_table(['项目', '内容'], meta, [1.4, 4.6])
doc.add_paragraph()

# ============================================================
# 一、执行摘要
# ============================================================
doc.add_heading('一、执行摘要', level=1)

add_para(
    f'本报告针对 B站视频 {BVID} 的评论区和弹幕数据进行系统性采集、清洗与分析。'
    f'B站前台显示该视频共有 {TOTAL_BILI:,} 条评论，通过异步爬虫实际获取 {CMT_CRAWLED:,} 条评论'
    f'（API 深度限制导致约 {TOTAL_BILI - CMT_CRAWLED:,} 条不可达），'
    f'以及 {DM_RAW_COUNT:,} 条原始弹幕（覆盖 {len(dm_daily)} 天）。'
    f'经七层规则清洗后，有效评论 {CMT_VALID:,} 条（保留率 {CMT_RET:.1f}%）。'
    f'弹幕经三道清洗+文本去重后，得 {DM_DEDUP:,} 种唯一文本（原始保留率 {DEDUP_RATE:.1f}%），'
    f'重复总次数 {DM_REPEAT_TOT:,}，最高单条重复 {DM_REPEAT_MAX} 次。'
)

# ============================================================
# 二、视频概览
# ============================================================
doc.add_heading('二、视频基本信息', level=1)

add_para(f'视频时长：约 {VIDEO_DUR:.0f} 秒（{VIDEO_DUR/60:.1f} 分钟）', space_after=2)
add_para(f'弹幕覆盖天数：{len(dm_daily)} 天', space_after=2)
add_para(f'弹幕峰值日：{max(dm_daily, key=dm_daily.get)}（{dm_daily[max(dm_daily, key=dm_daily.get)]:,} 条）', space_after=2)
add_para(f'评论峰值日：{max(cmt_daily, key=cmt_daily.get)}（{cmt_daily[max(cmt_daily, key=cmt_daily.get)]:,} 条）', space_after=10)

# ============================================================
# 三、数据漏斗
# ============================================================
doc.add_heading('三、数据漏斗', level=1)

doc.add_heading('评论数据漏斗', level=3)
add_para('七层规则逐级过滤，最终保留率 ' + f'{CMT_RET:.1f}%' + '。', space_after=6)
cmt_funnel = [
    ('B站前台显示总评论数', f'{TOTAL_BILI:,}', '100%'),
    ('API不可达（删除/隐藏/限制）', f'-{TOTAL_BILI - CMT_CRAWLED:,}', f'-{(TOTAL_BILI-CMT_CRAWLED)/max(TOTAL_BILI,1)*100:.1f}%'),
    ('实际爬取', f'{CMT_CRAWLED:,}', f'{CMT_CRAWLED/max(TOTAL_BILI,1)*100:.1f}%'),
    ('纯表情/颜文字', f'-{cmt_emoji:,}', f'-{cmt_emoji/max(CMT_CRAWLED,1)*100:.1f}%'),
    ('广告/营销', f'-{cmt_ad:,}', f'-{cmt_ad/max(CMT_CRAWLED,1)*100:.1f}%'),
    ('其他无效（短/灌水/符号等）', f'-{cmt_other:,}', f'-{cmt_other/max(CMT_CRAWLED,1)*100:.1f}%'),
    ('有效评论', f'{CMT_VALID:,}', f'{CMT_RET:.1f}%'),
]
make_table(['阶段', '数量', '占比'], cmt_funnel, [2.8, 1.3, 0.9])
doc.add_paragraph()

doc.add_heading('弹幕数据漏斗', level=3)
add_para('三道轻量清洗 + 文本去重。重复次数是弹幕的热度标尺，类比评论点赞。', space_after=6)
dm_funnel = [
    ('原始弹幕爬取量', f'{DM_RAW_COUNT:,}', '100%'),
    ('空弹幕剔除', f'-{dm_stage.get("empty",0):,}', f'-{dm_stage.get("empty",0)/max(DM_RAW_COUNT,1)*100:.1f}%'),
    ('单字符剔除（?/！等除外）', f'-{dm_stage.get("single",0):,}', f'-{dm_stage.get("single",0)/max(DM_RAW_COUNT,1)*100:.1f}%'),
    ('块状装饰符剔除', f'-{dm_stage.get("block",0):,}', f'-{dm_stage.get("block",0)/max(DM_RAW_COUNT,1)*100:.1f}%'),
    ('清洗后', f'{DM_CLEAN:,}', f'{DM_CLEAN/max(DM_RAW_COUNT,1)*100:.1f}%'),
    ('去重后（唯一文本数）', f'{DM_DEDUP:,}', f'{DEDUP_RATE:.1f}%'),
    ('最高单条重复', f'{DM_REPEAT_MAX}x', ''),
]
make_table(['阶段', '数量', '占比/备注'], dm_funnel, [2.8, 1.3, 0.9])
doc.add_paragraph()

# ============================================================
# 四、数据可视化
# ============================================================
doc.add_heading('四、数据可视化', level=1)

add_para('图1：评论（蓝色）与弹幕（粉色）每日发布量对比。', space_after=6)
add_img(IMG_DIR / 'chart1_daily.png', 6.3)

add_para('图2：弹幕重复次数排名（前200）。重复次数是弹幕的热度指标——越高代表越多观众发送了相同文本，反映集体情绪和高频梗。', space_after=6)
add_img(IMG_DIR / 'chart2_dm_freq.png', 5.5)

add_para('图3：弹幕首次出现位置在视频时间轴上的密度分布。峰值区域对应视频中最能激发观众表达欲的片段。', space_after=6)
add_img(IMG_DIR / 'chart3_dm_density.png', 6.3)

add_para('图4：评论点赞数分布（对数坐标）。绝大多数评论点赞不足10，少量高赞评论构成长尾。', space_after=6)
add_img(IMG_DIR / 'chart4_likes.png', 5.5)

add_para('图5：最活跃的15位评论用户及发评数量。', space_after=6)
add_img(IMG_DIR / 'chart5_top_users.png', 5.5)

add_para('图6：弹幕高频词Top20（按重复次数加权）。揭示观众讨论的情感焦点与核心议题。', space_after=6)
add_img(IMG_DIR / 'chart6_dm_words.png', 5.5)

add_para('图7：内容体量对比——有效评论 vs 弹幕去重前 vs 弹幕去重后。', space_after=6)
add_img(IMG_DIR / 'chart7_compare.png', 4.5)

# ============================================================
# 五、统计汇总
# ============================================================
doc.add_heading('五、统计汇总', level=1)

doc.add_heading('评论统计', level=3)
cmt_stats = [
    ('有效评论总数', f'{CMT_VALID:,} 条'),
    ('主评论', f'{len(comments_main):,} 条'),
    ('子回复', f'{len(comments_sub):,} 条'),
    ('评论保留率', f'{CMT_RET:.1f}%'),
    ('最高点赞', f'{max(cmt_likes) if cmt_likes else 0:,}'),
    ('平均点赞', f'{sum(cmt_likes)/max(len(cmt_likes),1):.0f}'),
    ('平均字数', f'{sum(cmt_lengths)/max(len(cmt_lengths),1):.0f} 字'),
    ('高赞评论（≥100赞）', f'{sum(1 for v in cmt_likes if v>=100):,} 条'),
    ('最高评论用户', f'{cmt_top_users.most_common(1)[0][0]}（{cmt_top_users.most_common(1)[0][1]}条）'),
]
make_table(['指标', '数值'], cmt_stats, [2.6, 3.4])
doc.add_paragraph()

doc.add_heading('弹幕统计（去重后）', level=3)
dm_stats = [
    ('原始弹幕爬取量', f'{DM_RAW_COUNT:,} 条'),
    ('去重后唯一文本数', f'{DM_DEDUP:,} 条'),
    ('去重率', f'{DEDUP_RATE:.1f}%'),
    ('重复总次数', f'{DM_REPEAT_TOT:,}'),
    ('平均每条重复', f'{DM_REPEAT_TOT/max(DM_DEDUP,1):.1f} 次'),
    ('最高单条重复', f'{DM_REPEAT_MAX} 次'),
    ('最高重复弹幕内容', f'{danmaku_dedup[0]["text"][:30]}...' if danmaku_dedup else ''),
    ('视频时长', f'{VIDEO_DUR:.0f} 秒（{VIDEO_DUR/60:.1f} 分钟）'),
    ('弹幕密度（去重后）', f'{DM_DEDUP/max(VIDEO_DUR,1):.1f} 条/秒'),
    ('覆盖天数', f'{len(dm_daily)} 天'),
]
make_table(['指标', '数值'], dm_stats, [2.6, 3.4])
doc.add_paragraph()

# ============================================================
# 六、清洗规则
# ============================================================
doc.add_heading('六、清洗规则', level=1)

doc.add_heading('评论清洗规则（七层）', level=3)
add_para('以下七类规则按优先级依次应用，命中任一规则即标记为无效并剔除。', space_after=6)
rules = [
    ('1', '纯表情/颜文字', '去除 Emoji 与颜文字后，无实质中文（<2字）且无英文单词'),
    ('2', '广告/营销', '匹配网址、手机号、微信号、QQ号、扫码、兼职刷单等'),
    ('3', '字数 ≤ 2', '评论内容不足 3 个字符'),
    ('4', '灌水词', '命中预定义常见无意义回复白名单（打卡/第一/三连等）'),
    ('5', '抢楼', '纯数字加"楼"字的占位评论'),
    ('6', '重复单字', '同一字符重复 3 次以上且全文无其他内容'),
    ('7', '纯符号', '全文由数字、标点、符号组成，无实质文字'),
]
make_table(['优先级', '规则', '判定逻辑'], rules, [0.5, 1.2, 4.3])
doc.add_paragraph()

doc.add_heading('弹幕清洗规则', level=3)
add_para('弹幕天然短小、情绪化，不宜照搬评论的重度清洗。采用"三道轻量清洗+文本去重"策略：', space_after=6)

doc.add_heading('第一阶段：噪声剔除', level=3)
add_para('三道确定性规则依次应用。规则2设有例外：? ？ ! ！ 作为情绪信号予以保留。', space_after=4)
dm_rules = [
    ('1', '空弹幕', '文本去除空白后为空'),
    ('2', '单字符剔除', '''文本长度为1，但 ? ？ ! ！ 作为情绪信号保留。其余单字符（如 6、.）剔除。'''),
    ('3', '块状装饰符', '全文由块状装饰符（Unicode U+2580-U+259F、U+25A0-U+25FF：▀▁▂▃▄▅▆▇█▉▊▋▌▍▎▏▐░▒▓等）及纯符号组成，无分析价值。'),
]
make_table(['优先级', '规则', '判定逻辑'], dm_rules, [0.5, 1.2, 4.3])
doc.add_paragraph()

doc.add_heading('第一阶段结果', level=3)
dm_clean = [
    ('原始弹幕', f'{DM_RAW_COUNT:,}', '100%'),
    ('空弹幕剔除', f'-{dm_stage.get("empty",0):,}', f'-{dm_stage.get("empty",0)/max(DM_RAW_COUNT,1)*100:.1f}%'),
    ('单字符剔除', f'-{dm_stage.get("single",0):,}', f'-{dm_stage.get("single",0)/max(DM_RAW_COUNT,1)*100:.1f}%'),
    ('块装饰符剔除', f'-{dm_stage.get("block",0):,}', f'-{dm_stage.get("block",0)/max(DM_RAW_COUNT,1)*100:.1f}%'),
    ('清洗后', f'{DM_CLEAN:,}', f'{DM_CLEAN/max(DM_RAW_COUNT,1)*100:.1f}%'),
]
make_table(['阶段', '数量', '占比'], dm_clean, [2.8, 1.3, 0.9])
doc.add_paragraph()

doc.add_heading('第二阶段：文本去重', level=3)
add_para('清洗后的弹幕按文本精确匹配分组。对每个唯一文本：', space_after=4)
add_para('• 重复次数：该文本总共发送了多少次', space_after=1)
add_para('• 首次出现：最早出现的视频时间点和发送时间', space_after=1)
add_para('• 聚合信息：最常见类型、颜色、弹幕池、活跃天数', space_after=1)
add_para('• 排序：按重复次数降序（弹幕世界的点赞）', space_after=6)

add_para('例如：若某弹幕文本在5天内出现847次，则合并为一行，携带重复次数=847、首次出现时间及聚合元数据。将原始冗余数据转化为有意义的流行度信号。', space_after=6)

dm_dedup = [
    ('清洗后', f'{DM_CLEAN:,}', '100%'),
    ('去重后唯一文本数', f'{DM_DEDUP:,}', f'{DM_DEDUP/max(DM_CLEAN,1)*100:.1f}%'),
    ('去重率', f'{DEDUP_RATE:.1f}%', '唯一文本 / 原始'),
    ('重复总次数', f'{DM_REPEAT_TOT:,}', f'平均每条 {DM_REPEAT_TOT/max(DM_DEDUP,1):.1f} 次'),
    ('最高单条重复', f'{DM_REPEAT_MAX} 次', ''),
]
make_table(['阶段', '数量', '备注'], dm_dedup, [2.8, 1.3, 1.9])

add_para('设计理念：评论的点赞是外部指标，而弹幕无原生互动信号。重复本身就是弹幕的流行度——越多人发出同样的文字，集体共鸣越强烈。去重将这一隐性信号变为显性、可排序的分析指标。', space_after=10)

# ============================================================
# 七、文件索引
# ============================================================
doc.add_heading('七、输出文件索引', level=1)

files = [
    ('comments.json', 'JSON', '原始爬取的评论数据'),
    ('comments_sorted.json', 'JSON', '清洗后按点赞排序的有效评论'),
    ('comments_sorted.xlsx', 'Excel', '清洗后的评论 Excel，含筛选与高亮'),
    (f'danmaku_{BVID}.json', 'JSON', '原始爬取的弹幕数据'),
    ('danmaku_dedup.json', 'JSON', '去重后按重复次数排序的弹幕'),
    ('danmaku_dedup.xlsx', 'Excel', '去重后的弹幕 Excel，高重复高亮'),
    ('B站视频综合分析报告.docx', 'Word', '本报告'),
]
make_table(['文件名', '格式', '说明'], files, [2.4, 0.7, 2.9])

# ============================================================
print('[6/7] 保存...')
doc.save(str(OUT_DOCX))
print(f'\n✅ 报告已生成: {OUT_DOCX}')
