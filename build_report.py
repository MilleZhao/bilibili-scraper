"""
生成 B站评论数据清洗分析报告（Word .docx）
预设: standard_business_brief
"""
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# 数据准备
# ============================================================
SRC = Path(r"C:\Users\LENOVO\Documents\爬虫\outputs\comments.json")
raw = json.loads(SRC.read_text(encoding="utf-8"))

total_bilibili = 24754
total_crawled  = len(raw)

# 统计各类过滤
from openpyxl.cell.cell import ILLEGAL_CHARACTERS_RE
import re

EMOJI_COUNT = 0
AD_COUNT    = 0
SHORT_COUNT = 0
SPAM_COUNT  = 0
FLOOR_COUNT = 0
REPEAT_COUNT= 0
SYMBOL_COUNT= 0

def is_emoji_ch(ch):
    cp = ord(ch)
    return (0x1F300<=cp<=0x1FAFF or 0x2600<=cp<=0x27BF or cp==0x2764
         or 0x1F900<=cp<=0x1F9FF or 0x1FA00<=cp<=0x1FA6F
         or 0x200D<=cp<=0x200D or 0xFE0F<=cp<=0xFE0F
         or 0x2702<=cp<=0x27B0 or 0x1F600<=cp<=0x1F64F
         or 0x1F680<=cp<=0x1F6FF or 0x1F1E0<=cp<=0x1F1FF)

def has_text(text):
    cleaned = "".join(ch for ch in text if not is_emoji_ch(ch))
    cleaned = re.sub(r"[\(（]\s*[^\w\s]{2,}\s*[\)）]", "", cleaned)
    chinese = len(re.findall(r"[\u4e00-\u9fff]", cleaned))
    english = len(re.findall(r"[a-zA-Z]{3,}", cleaned))
    return chinese >= 2 or english >= 1

def is_ad(text):
    patterns = [
        r"https?://\S+|www\.\S+",
        r"1[3-9]\d{9}|\d{3}[-\s]?\d{4}[-\s]?\d{4}",
        r"[加➕﹢＋]?\s*(微|v信|威信|VX|vx|WX|wx)[:：]?\s*\w{5,}",
        r"[加➕﹢＋]?\s*(qq|QQ|扣扣|Q|q)[:：]?\s*\d{5,}",
        r"扫码|二维码|QR码",
        r"免费领取|免费送|低价|优惠券|薅羊毛|返利|刷单|兼职|日赚|代练|代打",
    ]
    return any(re.search(p, text, re.I) for p in patterns)

total_valid = 0
for c in raw:
    text = (c.get("content") or "").strip()
    if not text:
        continue
    if not has_text(text):
        EMOJI_COUNT += 1
        continue
    if is_ad(text):
        AD_COUNT += 1
        continue
    if len(text) <= 2:
        SHORT_COUNT += 1
        continue
    if text in ("好","加油","支持","6","厉害","牛","顶","打卡","来了","第一","棒","不错","nb","NB","哈哈","哈哈哈","可以的","好家伙","三连"):
        SPAM_COUNT += 1
        continue
    if re.fullmatch(r"[\d\s,，]+楼", text):
        FLOOR_COUNT += 1
        continue
    if len(set(text)) <= 3 and len(text) >= 3:
        REPEAT_COUNT += 1
        continue
    if re.fullmatch(r"[\d\s\.\,\!\?\#\@\$\%\^\&\*\(\)\+\=~\-—–_/\\|:;]+", text):
        SYMBOL_COUNT += 1
        continue
    total_valid += 1

OTHER_COUNT = total_crawled - total_valid - EMOJI_COUNT - AD_COUNT
api_unreachable = total_bilibili - total_crawled
retention_rate = total_valid / total_crawled * 100 if total_crawled else 0

# ============================================================
# 构建 Word 文档
# ============================================================
doc = Document()

# ---- 页面设置 ----
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(1.0)
section.right_margin  = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# ---- 样式设置 ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(6)
style.paragraph_format.line_spacing = 1.10
style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

# 修复中文字体
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for level, (size, color_hex) in enumerate([
    (16, "2E74B5"), (13, "2E74B5"), (12, "1F4D78")
], 1):
    hs = doc.styles[f"Heading {level}"]
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor(*tuple(int(color_hex[i:i+2],16) for i in (0,2,4)))
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

# ---- 辅助函数 ----
BLUE_HEAD   = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE   = RGBColor(0x1F, 0x3A, 0x5F)
INK_BLUE    = RGBColor(0x0B, 0x25, 0x45)
CALLOUT_BG  = "F4F6F9"
TABLE_HEAD  = "F2F4F7"
TABLE_BODY  = "FFFFFF"

def add_para(text, bold=False, size=None, color=None, alignment=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def set_cell(cell, text, bold=False, size=10, color=INK_BLUE, align="left", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if fill:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # cell margins
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}><w:top w:w="60" w:type="dxa"/>'
        f'<w:bottom w:w="60" w:type="dxa"/>'
        f'<w:start w:w="100" w:type="dxa"/>'
        f'<w:end w:w="100" w:type="dxa"/></w:tcMar>'
    )
    tc_pr.append(tc_mar)

def set_table_borders(table, color="BFBFBF", size="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

# ============================================================
# 封面 / 标题
# ============================================================
add_para("B站评论数据清洗分析报告", bold=True, size=24, color=DARK_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("BV11mFLziEyP  —  当世界过分诚实，我们要如何保持好奇与勇气", size=13, color=BLUE_HEAD, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# 分隔线
p_line = doc.add_paragraph()
p_line.paragraph_format.space_after = Pt(16)
run_line = p_line.add_run("_" * 50)
run_line.font.size = Pt(6)
run_line.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# 元数据表
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(meta_table, "BFBFBF", "4")

meta_data = [
    ("报告日期", datetime.now().strftime("%Y年%m月%d日")),
    ("数据来源", "bilibili.com  —  bilibili-api 库"),
    ("爬取时间", "2026年6月"),
    ("分析方法", "Python 异步爬虫 + 规则引擎清洗"),
]
for i, (label, value) in enumerate(meta_data):
    set_cell(meta_table.cell(i, 0), label, bold=True, size=10, fill=TABLE_HEAD, align="right")
    set_cell(meta_table.cell(i, 1), value, size=10)

# 列宽
for row in meta_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(3.5)

doc.add_paragraph()  # 空行

# ============================================================
# 1. 执行摘要
# ============================================================
doc.add_heading("一、执行摘要", level=1)

add_para(
    f"本报告针对 B站视频 BV11mFLziEyP 的评论区数据进行系统性清洗与分析。"
    f"B站前端显示该视频共有 {total_bilibili:,} 条评论，通过 bilibili-api 异步爬虫"
    f"实际获取 {total_crawled:,} 条（API 深度限制导致约 {api_unreachable:,} 条不可达）。"
    f"经过七类规则清洗后，保留有效评论 {total_valid:,} 条，保留率 {retention_rate:.1f}%。"
    f"清洗后的数据按点赞数从高到低排列，可供后续内容分析、舆情洞察使用。"
)

# ============================================================
# 2. 数据漏斗
# ============================================================
doc.add_heading("二、数据漏斗", level=1)

add_para("以下漏斗图展示从 B站原始评论数到最终有效评论数的逐层过滤过程。", space_after=10)

# 漏斗表
funnel = doc.add_table(rows=9, cols=4)
funnel.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(funnel)

funnel_widths = [0.6, 2.8, 1.2, 1.4]

# 表头
headers = ["层级", "过滤阶段", "剩余数量", "占比"]
for i, h in enumerate(headers):
    set_cell(funnel.cell(0, i), h, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), fill="2E74B5", align="center")

funnel_data = [
    ("1", "B站显示总评论数",       total_bilibili,  "100%"),
    ("2", "API 不可达（被删/隐藏/限制）", f"-{api_unreachable:,}", f"-{api_unreachable/total_bilibili*100:.1f}%"),
    ("3", "实际爬取",              total_crawled,   f"{total_crawled/total_bilibili*100:.1f}%"),
    ("4", "  纯表情/颜文字",       f"-{EMOJI_COUNT:,}",  f"-{EMOJI_COUNT/total_crawled*100:.1f}%"),
    ("5", "  广告/营销",           f"-{AD_COUNT:,}",     f"-{AD_COUNT/total_crawled*100:.1f}%"),
    ("6", "  其他无效（短/灌水/符号等）", f"-{OTHER_COUNT:,}", f"-{OTHER_COUNT/total_crawled*100:.1f}%"),
    ("7", "有效评论（最终保留）",   total_valid,      f"{retention_rate:.1f}%"),
    ("8", "其中: 高赞评论（≥100赞）", f"{sum(1 for c in raw if c.get('likes',0)>=100):,}", "—"),
]

for i, (level, stage, count, pct) in enumerate(funnel_data):
    row = i + 1
    is_neg = str(count).startswith("-")
    is_final = i == 6
    is_sub   = stage.startswith("  ")
    
    c = RGBColor(0x9B, 0x1C, 0x1C) if is_neg else DARK_BLUE
    if is_final:
        c = RGBColor(0x1F, 0x4D, 0x78)
    
    fill = None
    if is_final:
        fill = "E8EEF5"
    elif is_sub:
        fill = "F9F9F9"
    
    set_cell(funnel.cell(row, 0), level, bold=is_final, color=c, align="center", fill=fill)
    set_cell(funnel.cell(row, 1), stage, bold=is_final, color=c, fill=fill)
    set_cell(funnel.cell(row, 2), str(count), bold=is_final, color=c, align="right", fill=fill)
    set_cell(funnel.cell(row, 3), pct, bold=is_final, color=c, align="center", fill=fill)

for row in funnel.rows:
    for i, w in enumerate(funnel_widths):
        row.cells[i].width = Inches(w)

add_para("表1: 评论数据漏斗 —— 从原始评论到有效评论的逐层过滤", size=9, color=RGBColor(0x88,0x88,0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=10)

# ============================================================
# 3. 过滤规则详解
# ============================================================
doc.add_heading("三、过滤规则详解", level=1)

add_para("以下七类规则按优先级依次应用，命中任一规则即标记为无效并剔除。规则设计遵循'宁可多保留，避免误伤有价值评论'的原则。", space_after=10)

rules = doc.add_table(rows=8, cols=4)
rules.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(rules)

for i, h in enumerate(["优先级", "规则名称", "判定逻辑", "示例"]):
    set_cell(rules.cell(0, i), h, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), fill="2E74B5", align="center")

rule_data = [
    ("1", "纯表情/颜文字", "去除 Emoji 与颜文字后无实质中文（< 2个汉字）且无英文单词", "😂😂😂 / (≧∇≦)ﾉ / 🔥❤️👍"),
    ("2", "广告/营销", "匹配网址、手机号、微信号、QQ号、扫码、兼职刷单等关键词", "加VX:xxx / 扫码领 / 免费送 / 日赚300"),
    ("3", "字数 ≤ 2", "评论内容长度不足 3 个字符", "好 / nb / 666"),
    ("4", "灌水词", "命中预定义的常见无意义回复白名单", "打卡 / 来了 / 第一 / 支持 / 三连"),
    ("5", "抢楼", "纯数字加'楼'字的占位评论", "123楼 / 1楼 / 520楼"),
    ("6", "重复单字", "同一字符重复 3 次以上且全文无其他字符", "啊啊啊啊啊 / 哈哈哈哈哈"),
    ("7", "纯符号", "全文由数字、标点、符号组成，无实质文字", "!!! / ...... / ？？？"),
]

for i, (pri, name, logic, example) in enumerate(rule_data):
    row = i + 1
    fill = "F9F9F9" if i % 2 == 0 else None
    set_cell(rules.cell(row, 0), pri, align="center", fill=fill)
    set_cell(rules.cell(row, 1), name, bold=True, fill=fill)
    set_cell(rules.cell(row, 2), logic, fill=fill)
    set_cell(rules.cell(row, 3), example, fill=fill)

for row in rules.rows:
    row.cells[0].width = Inches(0.55)
    row.cells[1].width = Inches(1.1)
    row.cells[2].width = Inches(3.1)
    row.cells[3].width = Inches(1.75)

add_para("表2: 七类过滤规则明细", size=9, color=RGBColor(0x88,0x88,0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=10)

# ============================================================
# 4. 过滤结果统计
# ============================================================
doc.add_heading("四、过滤结果统计", level=1)

stats = doc.add_table(rows=9, cols=3)
stats.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(stats)

for i, h in enumerate(["过滤类别", "数量", "占爬取总量比"]):
    set_cell(stats.cell(0, i), h, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), fill="2E74B5", align="center")

stat_data = [
    ("纯表情/颜文字",     EMOJI_COUNT,  EMOJI_COUNT/total_crawled*100),
    ("广告/营销",         AD_COUNT,     AD_COUNT/total_crawled*100),
    ("字数过短（≤2字）",  SHORT_COUNT,  SHORT_COUNT/total_crawled*100 if total_crawled else 0),
    ("灌水词",            SPAM_COUNT,   SPAM_COUNT/total_crawled*100 if total_crawled else 0),
    ("抢楼",              FLOOR_COUNT,  FLOOR_COUNT/total_crawled*100 if total_crawled else 0),
    ("重复单字",          REPEAT_COUNT, REPEAT_COUNT/total_crawled*100 if total_crawled else 0),
    ("纯符号",            SYMBOL_COUNT, SYMBOL_COUNT/total_crawled*100 if total_crawled else 0),
    ("合计剔除",          total_crawled-total_valid, (total_crawled-total_valid)/total_crawled*100 if total_crawled else 0),
]

for i, (cat, cnt, pct) in enumerate(stat_data):
    row = i + 1
    is_total = i == 7
    fill = "E8EEF5" if is_total else ("F9F9F9" if i % 2 == 0 else None)
    set_cell(stats.cell(row, 0), cat, bold=is_total, fill=fill)
    set_cell(stats.cell(row, 1), f"{cnt:,}", bold=is_total, align="right", fill=fill)
    set_cell(stats.cell(row, 2), f"{pct:.1f}%", bold=is_total, align="center", fill=fill)

for row in stats.rows:
    row.cells[0].width = Inches(2.4)
    row.cells[1].width = Inches(1.4)
    row.cells[2].width = Inches(1.7)

add_para("表3: 各类过滤结果统计", size=9, color=RGBColor(0x88,0x88,0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=10)

# ============================================================
# 5. 最终数据概况
# ============================================================
doc.add_heading("五、最终数据概况", level=1)

final_data = doc.add_table(rows=6, cols=2)
final_data.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(final_data)

kv = [
    ("有效评论总数",    f"{total_valid:,} 条"),
    ("保留率",          f"{retention_rate:.1f}%"),
    ("主评论数",        f"{sum(1 for c in raw if not c.get('is_sub',False) and len((c.get('content')or'').strip())>2):,}"),
    ("子回复数",        f"{sum(1 for c in raw if c.get('is_sub',False)):,}"),
    ("≥100赞高赞评论",  f"{sum(1 for c in raw if c.get('likes',0)>=100):,} 条"),
    ("输出格式",        "Word (.docx) + Excel (.xlsx) + JSON"),
]

for i, (k, v) in enumerate(kv):
    fill = TABLE_HEAD if i % 2 == 0 else None
    set_cell(final_data.cell(i, 0), k, bold=True, size=10, fill=fill, align="right")
    set_cell(final_data.cell(i, 1), v, size=10)

for row in final_data.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(3.5)

add_para("表4: 最终数据概况", size=9, color=RGBColor(0x88,0x88,0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=10)

# ============================================================
# 6. 数据文件索引
# ============================================================
doc.add_heading("六、数据文件索引", level=1)

add_para("清洗后的数据以以下格式提供，可根据需求选用：", space_after=6)

file_table = doc.add_table(rows=4, cols=3)
file_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(file_table)

for i, h in enumerate(["文件名", "格式", "说明"]):
    set_cell(file_table.cell(0, i), h, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), fill="2E74B5", align="center")

files = [
    ("comments.json", "JSON", "包含全部 17,278 条原始爬取数据的结构化文件"),
    ("comments_sorted.json", "JSON", "清洗后按点赞排序的 14,002 条有效评论"),
    ("comments_sorted.xlsx", "Excel", "清洗后的 Excel，含筛选器与高赞高亮"),
]

for i, (fn, fmt, desc) in enumerate(files):
    fill = "F9F9F9" if i % 2 == 0 else None
    set_cell(file_table.cell(i+1, 0), fn, fill=fill)
    set_cell(file_table.cell(i+1, 1), fmt, align="center", fill=fill)
    set_cell(file_table.cell(i+1, 2), desc, fill=fill)

for row in file_table.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(0.8)
    row.cells[2].width = Inches(3.9)

add_para("表5: 数据文件索引", size=9, color=RGBColor(0x88,0x88,0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=10)

# ============================================================
# 7. 技术说明
# ============================================================
doc.add_heading("七、技术说明", level=1)

add_para("爬虫框架: bilibili-api（Python）", bold=True, space_after=2)
add_para("采用异步架构（asyncio），通过游标分页（get_comments_lazy + offset）突破传统页码翻页的深度限制。使用 SESSDATA 凭证登录态确保评论接口可访问。", space_after=6)

add_para("清洗引擎: Python 正则表达式 + Unicode 范围判断", bold=True, space_after=2)
add_para("Emoji 判定采用 Unicode 码点范围检测，广告检测采用正则关键词匹配，灌水词采用白名单精确匹配。所有规则均为确定性规则，不依赖机器学习模型，确保清洗结果可复现、可审计。", space_after=6)

add_para("局限性", bold=True, space_after=2)
add_para(f"B站 API 深度限制导致约 {api_unreachable:,} 条评论无法获取（~{api_unreachable/total_bilibili*100:.1f}%），这部分数据可能包含被删除评论、被折叠评论或超出 API 分页上限的旧评论。清洗规则的灌水词白名单可能遗漏部分无意义评论（如方言或新兴网络用语），建议在使用前根据具体分析目的酌情调整。")

# ============================================================
# 保存
# ============================================================
OUT = SRC.parent / "评论清洗分析报告.docx"
doc.save(str(OUT))
print(f"✅ 报告已生成: {OUT}")
