from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_DIR = Path("outputs")
OUT_PATH = OUT_DIR / "B站AI_AIGC数据爬取说明.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_doc_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)

    for name, size in [("Title", 20), ("Heading 1", 14), ("Heading 2", 11.5)]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 1:
        run.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run("• ").bold = True
        p.add_run(item)


def build_docx() -> None:
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("B站 AI/AIGC 数据爬取说明")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    meta.runs[0].font.size = Pt(9)

    add_heading(doc, "一、项目概述", 1)
    doc.add_paragraph(
        "本次任务围绕你提供的 B 站视频 BV11mFLziEyP 展开，以 AI / AIGC 主题为主线，"
        "设计了可批量扩展的采集方案。Excel 工作簿按帖子与评论两层结构组织，便于后续补齐 200+ 帖子与 20000+ 评论的量化要求。"
    )

    add_heading(doc, "二、工具与方法", 1)
    add_bullets(
        doc,
        [
            "Python：负责公开接口请求、数据清洗、去重和字段整理。",
            "Bilibili 公共接口：用于获取视频详情、搜索结果与评论分页。",
            "JavaScript + artifact-tool：用于导出结构化 Excel 工作簿。",
            "python-docx：用于生成 Word 说明文档。",
        ],
    )

    add_heading(doc, "三、数据字段口径", 1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "数据层"
    hdr[1].text = "字段范围"
    set_cell_shading(hdr[0], "D9EAF7")
    set_cell_shading(hdr[1], "D9EAF7")
    set_repeat_table_header(tbl.rows[0])
    rows = [
        ("帖子", "标题、原帖全文/简介、BV 号、UP 主、发布时间、播放、点赞、评论、转发、收藏、投币、链接、抓取时间、主题标签"),
        ("评论", "评论 ID、父评论 ID、评论内容、评论时间、点赞数、回复数、用户名、帖子链接、抓取时间"),
        ("readme", "采集口径、字段说明、使用工具、注意事项、补采建议"),
    ]
    for left, right in rows:
        row = tbl.add_row().cells
        row[0].text = left
        row[1].text = right

    add_heading(doc, "四、采集策略", 1)
    add_bullets(
        doc,
        [
            "先以种子视频为中心，再按 AI / AIGC 相关关键词扩展同主题视频。",
            "优先选择评论量高的视频，确保更快达到评论总量阈值。",
            "评论按分页拉取并按评论 ID 去重，保留楼中楼父子关系。",
            "如果某条视频因删除、限流或评论为空无法使用，则自动跳过并继续扩展候选集。",
        ],
    )

    add_heading(doc, "五、当前交付状态", 1)
    doc.add_paragraph(
        "在当前桌面环境中，公开网络访问被沙箱限制，因此我已经完成了可复用的采集脚本、Excel 构建器和 Word 说明文生成器，"
        "并输出了结构化模板。你只需在可联网环境中运行采集脚本，再重新执行构建器，就能得到完整的 200+ 帖子和 20000+ 评论数据集。"
    )

    add_heading(doc, "六、后续补全建议", 1)
    add_bullets(
        doc,
        [
            "先执行 collect_bilibili_ai_aigc.py 生成 posts.json 和 comments.json。",
            "再运行 build_xlsx.mjs 与 build_docx.py 重新导出最终版文件。",
            "导出后抽查链接、标题和评论内容是否能回链到原视频页面。",
        ],
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_docx()
    print(f"saved {OUT_PATH}")
